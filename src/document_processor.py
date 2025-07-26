a"""
Document Processing Service for ESG file uploads and analysis
"""
import io
import os
import logging
from typing import Dict, Any, List, Optional, Tuple
from datetime import datetime
import base64

# Document processing imports
try:
    import PyPDF2
    from docx import Document
    import pandas as pd
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    from .agents.base_agent import get_llm_service
except ImportError:
    # Fallback for import issues
    def get_llm_service():
        from .agents.base_agent import MockLLMService
        return MockLLMService()


class DocumentProcessor:
    """Process uploaded documents for ESG analysis"""
    
    def __init__(self):
        self.logger = logging.getLogger("document_processor")
        self.llm_service = get_llm_service()
        self.supported_formats = ['.pdf', '.docx', '.txt', '.csv', '.xlsx']
        self.max_file_size = 10 * 1024 * 1024  # 10MB limit
    
    def is_supported_format(self, filename: str) -> bool:
        """Check if file format is supported"""
        _, ext = os.path.splitext(filename.lower())
        return ext in self.supported_formats
    
    def validate_file(self, file_content: bytes, filename: str) -> Tuple[bool, str]:
        """Validate uploaded file"""
        if len(file_content) > self.max_file_size:
            return False, f"File size exceeds {self.max_file_size / (1024*1024):.1f}MB limit"
        
        if not self.is_supported_format(filename):
            return False, f"Unsupported file format. Supported formats: {', '.join(self.supported_formats)}"
        
        return True, "File validation successful"
    
    async def process_document(self, file_content: bytes, filename: str, document_type: str = "sustainability_report") -> Dict[str, Any]:
        """Process uploaded document and extract ESG-relevant information"""
        
        # Validate file
        is_valid, message = self.validate_file(file_content, filename)
        if not is_valid:
            return {
                "success": False,
                "error": message,
                "timestamp": datetime.now().isoformat()
            }
        
        try:
            # Extract text from document
            extracted_text = await self._extract_text(file_content, filename)
            
            if not extracted_text:
                return {
                    "success": False,
                    "error": "Could not extract text from document",
                    "timestamp": datetime.now().isoformat()
                }
            
            # Analyze document with LLM
            analysis = await self.llm_service.analyze_uploaded_document(
                extracted_text, document_type
            )
            
            # Extract ESG metrics and data
            esg_data = self._extract_esg_data(extracted_text)
            
            return {
                "success": True,
                "filename": filename,
                "document_type": document_type,
                "file_size": len(file_content),
                "text_length": len(extracted_text),
                "extracted_text": extracted_text[:1000] + "..." if len(extracted_text) > 1000 else extracted_text,
                "llm_analysis": analysis,
                "esg_data": esg_data,
                "processing_timestamp": datetime.now().isoformat()
            }
            
        except Exception as e:
            self.logger.error(f"Error processing document {filename}: {str(e)}")
            return {
                "success": False,
                "error": f"Document processing failed: {str(e)}",
                "timestamp": datetime.now().isoformat()
            }
    
    async def _extract_text(self, file_content: bytes, filename: str) -> str:
        """Extract text from various file formats"""
        _, ext = os.path.splitext(filename.lower())
        
        try:
            if ext == '.pdf':
                return self._extract_pdf_text(file_content)
            elif ext == '.docx':
                return self._extract_docx_text(file_content)
            elif ext == '.txt':
                return file_content.decode('utf-8', errors='ignore')
            elif ext == '.csv':
                return self._extract_csv_text(file_content)
            elif ext == '.xlsx':
                return self._extract_excel_text(file_content)
            else:
                return file_content.decode('utf-8', errors='ignore')
                
        except Exception as e:
            self.logger.error(f"Error extracting text from {filename}: {str(e)}")
            return ""
    
    def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF file"""
        if not PDF_SUPPORT:
            return "PDF processing not available - install PyPDF2"
        
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return text.strip()
            
        except Exception as e:
            self.logger.error(f"Error extracting PDF text: {str(e)}")
            return ""
    
    def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            doc_file = io.BytesIO(file_content)
            doc = Document(doc_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
            
        except Exception as e:
            self.logger.error(f"Error extracting DOCX text: {str(e)}")
            return ""
    
    def _extract_csv_text(self, file_content: bytes) -> str:
        """Extract text from CSV file"""
        try:
            csv_file = io.StringIO(file_content.decode('utf-8', errors='ignore'))
            df = pd.read_csv(csv_file)
            
            # Convert DataFrame to text representation
            text = f"CSV Data Summary:\n"
            text += f"Columns: {', '.join(df.columns.tolist())}\n"
            text += f"Rows: {len(df)}\n\n"
            text += "Sample Data:\n"
            text += df.head(10).to_string()
            
            return text
            
        except Exception as e:
            self.logger.error(f"Error extracting CSV text: {str(e)}")
            return ""
    
    def _extract_excel_text(self, file_content: bytes) -> str:
        """Extract text from Excel file"""
        try:
            excel_file = io.BytesIO(file_content)
            df = pd.read_excel(excel_file, sheet_name=None)  # Read all sheets
            
            text = "Excel Data Summary:\n"
            
            for sheet_name, sheet_df in df.items():
                text += f"\nSheet: {sheet_name}\n"
                text += f"Columns: {', '.join(sheet_df.columns.tolist())}\n"
                text += f"Rows: {len(sheet_df)}\n"
                text += "Sample Data:\n"
                text += sheet_df.head(5).to_string()
                text += "\n" + "="*50 + "\n"
            
            return text
            
        except Exception as e:
            self.logger.error(f"Error extracting Excel text: {str(e)}")
            return ""
    
    def _extract_esg_data(self, text: str) -> Dict[str, Any]:
        """Extract ESG-specific data points from text"""
        esg_data = {
            "environmental_metrics": [],
            "social_metrics": [],
            "governance_metrics": [],
            "financial_metrics": [],
            "targets_and_goals": [],
            "certifications": []
        }
        
        text_lower = text.lower()
        
        # Environmental keywords
        env_keywords = [
            "carbon", "co2", "emissions", "energy", "renewable", "waste", "water",
            "biodiversity", "climate", "greenhouse gas", "scope 1", "scope 2", "scope 3"
        ]
        
        # Social keywords
        social_keywords = [
            "diversity", "inclusion", "employee", "safety", "training", "community",
            "human rights", "gender", "equality", "workplace"
        ]
        
        # Governance keywords
        gov_keywords = [
            "board", "governance", "ethics", "compliance", "risk", "transparency",
            "audit", "independence", "oversight"
        ]
        
        # Extract metrics based on keywords
        lines = text.split('\n')
        for line in lines:
            line_lower = line.lower()
            
            # Check for environmental metrics
            if any(keyword in line_lower for keyword in env_keywords):
                if any(char.isdigit() for char in line):  # Contains numbers
                    esg_data["environmental_metrics"].append(line.strip())
            
            # Check for social metrics
            elif any(keyword in line_lower for keyword in social_keywords):
                if any(char.isdigit() for char in line):
                    esg_data["social_metrics"].append(line.strip())
            
            # Check for governance metrics
            elif any(keyword in line_lower for keyword in gov_keywords):
                if any(char.isdigit() for char in line):
                    esg_data["governance_metrics"].append(line.strip())
            
            # Check for targets and goals
            if any(word in line_lower for word in ["target", "goal", "objective", "commitment"]):
                esg_data["targets_and_goals"].append(line.strip())
            
            # Check for certifications
            if any(word in line_lower for word in ["certified", "certification", "iso", "gri", "sasb"]):
                esg_data["certifications"].append(line.strip())
        
        # Limit results to avoid overwhelming output
        for key in esg_data:
            esg_data[key] = esg_data[key][:10]  # Limit to 10 items per category
        
        return esg_data
    
    def get_document_insights(self, processed_documents: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Generate insights from multiple processed documents"""
        if not processed_documents:
            return {"insights": "No documents processed"}
        
        insights = {
            "total_documents": len(processed_documents),
            "document_types": {},
            "common_themes": [],
            "data_gaps": [],
            "compliance_indicators": [],
            "recommendations": []
        }
        
        # Analyze document types
        for doc in processed_documents:
            doc_type = doc.get("document_type", "unknown")
            insights["document_types"][doc_type] = insights["document_types"].get(doc_type, 0) + 1
        
        # Extract common themes from LLM analyses
        all_analyses = []
        for doc in processed_documents:
            if doc.get("success") and doc.get("llm_analysis"):
                all_analyses.append(doc["llm_analysis"].get("analysis", ""))
        
        # Simple theme extraction (could be enhanced with NLP)
        common_words = {}
        for analysis in all_analyses:
            words = analysis.lower().split()
            for word in words:
                if len(word) > 4:  # Only consider longer words
                    common_words[word] = common_words.get(word, 0) + 1
        
        # Get top themes
        sorted_themes = sorted(common_words.items(), key=lambda x: x[1], reverse=True)
        insights["common_themes"] = [theme[0] for theme in sorted_themes[:10]]
        
        # Generate recommendations based on document analysis
        insights["recommendations"] = [
            "Ensure all ESG data is quantified and measurable",
            "Align reporting with EU CSRD requirements",
            "Implement third-party verification for key metrics",
            "Establish clear ESG governance structure",
            "Develop comprehensive stakeholder engagement plan"
        ]
        
        return insights
    
    def create_download_link(self, file_content: bytes, filename: str) -> str:
        """Create download link for processed file"""
        b64 = base64.b64encode(file_content).decode()
        return f'<a href="data:application/octet-stream;base64,{b64}" download="{filename}">Download {filename}</a>'


# Global instance
document_processor = DocumentProcessor()